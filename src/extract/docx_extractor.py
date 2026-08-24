from __future__ import annotations

import os

os.environ["TORCH_COMPILE_DISABLE"] = "1"
os.environ["PYTORCH_JIT"] = "0"
os.environ["OMP_NUM_THREADS"] = "1"

import importlib.util
import logging
from pathlib import Path

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.extract.base import (
    BaseExtractor,
    DocumentSection,
    ExtractedDocument,
    ExtractedImage,
    generate_document_id,
)

logger = logging.getLogger(__name__)
_HAS_DOCLING = importlib.util.find_spec("docling") is not None


class DOCXExtractor(BaseExtractor):
    def extract(self, file_path: Path, role: str, category: str) -> ExtractedDocument:
        if _HAS_DOCLING:
            docling_result = self._try_docling_extract(file_path, role, category)
            if docling_result:
                return docling_result

        return self._docx_extract(file_path, role, category)


    def _try_docling_extract(self, file_path: Path, role: str, category: str) -> ExtractedDocument | None:
        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()
            res = converter.convert(str(file_path))
            docling_doc = res.document
            markdown_content = docling_doc.export_to_markdown()

            doc_id = generate_document_id(category, file_path.name)
            title = file_path.stem.replace("_", " ").replace("-", " ")
            num_pages = len(docling_doc.pages) if hasattr(docling_doc, "pages") and docling_doc.pages else 1

            sections = [
                DocumentSection(
                    title=title,
                    level=1,
                    content=markdown_content,
                    section_type="text",
                )
            ]

            return ExtractedDocument(
                document_id=doc_id,
                title=title,
                source_file=file_path.name,
                source_path=f"{category}/{file_path.name}",
                document_type="docx",
                role=role,
                category=category,
                pages=num_pages,
                sections=sections,
                images=[],
                raw_text=markdown_content,
            )
        except BaseException as e:
            logger.warning(f"Docling extraction skipped/failed for DOCX {file_path.name}: {e}")
            return None

    def _docx_extract(self, file_path: Path, role: str, category: str) -> ExtractedDocument:
        doc = docx.Document(file_path)
        sections: list[DocumentSection] = []
        images: list[ExtractedImage] = []
        raw_text_parts: list[str] = []

        current_heading = "Document Content"
        current_level = 1
        current_paragraphs: list[str] = []

        def flush_section():
            nonlocal current_paragraphs
            if current_paragraphs:
                content = "\n\n".join(current_paragraphs).strip()
                if content:
                    sections.append(
                        DocumentSection(
                            title=current_heading,
                            level=current_level,
                            content=content,
                            section_type="text",
                        )
                    )
                current_paragraphs = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                p = Paragraph(element, doc)
                text = p.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)
                style_name = p.style.name.lower() if p.style else ""

                if style_name.startswith("heading"):
                    flush_section()
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 1
                    current_heading = text
                    current_level = level
                elif style_name.startswith("list") or p._element.xpath("./w:pPr/w:numPr"):
                    current_paragraphs.append(f"- {text}")
                else:
                    current_paragraphs.append(text)

            elif element.tag.endswith("tbl"):
                tbl = Table(element, doc)
                table_md = self._convert_table_to_markdown(tbl)
                if table_md:
                    raw_text_parts.append(table_md)
                    current_paragraphs.append(table_md)

        flush_section()

        img_idx = 1
        for rel in doc.part.rels.values():
            if not rel.is_external and "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    ext = image_part.content_type.split("/")[-1]
                    if ext == "jpeg":
                        ext = "jpg"
                    img_name = f"{file_path.stem}_img{img_idx}.{ext}"
                    images.append(
                        ExtractedImage(
                            filename=img_name,
                            image_bytes=image_bytes,
                        )
                    )
                    img_idx += 1
                except Exception as e:
                    logger.debug(f"Skipping relationship {rel}: {e}")

        doc_id = generate_document_id(category, file_path.name)
        title = file_path.stem.replace("_", " ").replace("-", " ")

        return ExtractedDocument(
            document_id=doc_id,
            title=title,
            source_file=file_path.name,
            source_path=f"{category}/{file_path.name}",
            document_type="docx",
            role=role,
            category=category,
            pages=1,
            sections=sections,
            images=images,
            raw_text="\n\n".join(raw_text_parts),
        )

    def _convert_table_to_markdown(self, table: Table) -> str:
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_data.append(row_cells)

        if not rows_data:
            return ""

        headers = rows_data[0]
        markdown_lines = [
            "| " + " | ".join(headers) + " |",
            "| " + " | ".join(["---"] * len(headers)) + " |",
        ]

        for row in rows_data[1:]:
            if len(row) < len(headers):
                row.extend([""] * (len(headers) - len(row)))
            elif len(row) > len(headers):
                row = row[: len(headers)]
            markdown_lines.append("| " + " | ".join(row) + " |")

        return "\n".join(markdown_lines)
