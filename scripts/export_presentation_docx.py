from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Set inner padding of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f"<w:tcMar {nsdecls('w')}>"
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f"</w:tcMar>"
    )
    tcPr.append(tcMar)


def create_document():
    doc = Document()

    # Page Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    COLOR_PRIMARY = RGBColor(15, 107, 58)  # VinFast Green #0F6B3A
    COLOR_SECONDARY = RGBColor(30, 111, 181)  # Deep Blue #1E6FB5
    COLOR_DARK = RGBColor(33, 37, 41)  # Text Charcoal #212529
    COLOR_MUTED = RGBColor(108, 117, 125)  # Gray #6C757D
    COLOR_RED = RGBColor(227, 84, 63)  # Alert Red #E3543F

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("DỰ ÁN AI THỰC CHIẾN VINUNI — TEAM THE SIGMOID\n")
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_SECONDARY

    r_title = p_title.add_run("KỊCH BẢN THUYẾT TRÌNH & DEMO USER FLOW END-TO-END\n")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    r_app = p_title.add_run("VF AI Onboarding & Operational Support Assistant")
    r_app.font.size = Pt(13)
    r_app.font.italic = True
    r_app.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()  # Spacing

    # Callout Summary Box
    table_summary = doc.add_table(rows=1, cols=1)
    table_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sum = table_summary.rows[0].cells[0]
    set_cell_background(c_sum, "F0FDF4")  # Light green tint
    set_cell_margins(c_sum, 180, 180, 200, 200)
    p_box = c_sum.paragraphs[0]
    r_box_title = p_box.add_run("📌 THÔNG TIN CHUNG VỀ BÀI BÁO CÁO\n")
    r_box_title.font.bold = True
    r_box_title.font.size = Pt(11)
    r_box_title.font.color.rgb = COLOR_PRIMARY

    p_box.add_run("• Đội ngũ thực hiện: ").bold = True
    p_box.add_run("Team The Sigmoid (AI Thực chiến VinUni Cohort 3)\n")
    p_box.add_run("• Đề tài: ").bold = True
    p_box.add_run("Xây dựng Trợ lý AI Hội nhập và Hỗ trợ Tác nghiệp Chuyên sâu cho Đại lý VinFast\n")
    p_box.add_run("• Thời lượng thuyết trình đề xuất: ").bold = True
    p_box.add_run("8 – 10 phút (3 phút Slide + 5 phút Live Demo + 2 phút Q&A)\n")
    p_box.add_run("• Kiến trúc cốt lõi: ").bold = True
    p_box.add_run(
        "Cloud-Native (MinIO S3, PostgreSQL), Hybrid Search (ChromaDB + BM25), Cross-Encoder Reranker, LangGraph Multi-Agent, Phân quyền RBAC."
    )

    doc.add_paragraph()

    # Section 1: Presentation Script
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("I. LỜI THOẠI THUYẾT TRÌNH CHI TIẾT (SLIDE BY SLIDE)")
    r_h1.font.color.rgb = COLOR_PRIMARY

    # Slide 1
    h2_1 = doc.add_heading(level=2)
    h2_1.add_run("Slide 1: Trang bìa & Mở đầu (Thời lượng: 30 giây)")
    p_s1_act = doc.add_paragraph()
    p_s1_act.add_run("🎬 Thao tác màn hình: ").bold = True
    p_s1_act.add_run("Chiếu Slide Trang bìa (Logo The Sigmoid & VinFast).")

    p_s1_talk = doc.add_paragraph()
    p_s1_talk.add_run("🎙️ Lời thoại người trình bày:\n").bold = True
    r_talk1 = p_s1_talk.add_run(
        '"Kính chào Ban Giám khảo, các Quý Mentor và toàn thể hội đồng đang có mặt trong buổi báo cáo ngày hôm nay.\n'
        "Chúng em là Team The Sigmoid. Hôm nay, nhóm rất tự hào được giới thiệu giải pháp: "
        "VF AI Onboarding & Operational Support — Nền tảng Trợ lý AI Hội nhập và Hỗ trợ Tác nghiệp Chuyên sâu "
        'dành riêng cho mạng lưới Đại lý VinFast trên toàn quốc."'
    )
    r_talk1.font.italic = True

    # Slide 2
    h2_2 = doc.add_heading(level=2)
    h2_2.add_run("Slide 2: Bối cảnh & 3 Thách thức Thực tế tại Đại lý (Thời lượng: 1 phút)")
    p_s2_act = doc.add_paragraph()
    p_s2_act.add_run("🎬 Thao tác màn hình: ").bold = True
    p_s2_act.add_run("Chuyển sang Slide 3 Nỗi đau lớn của nhân sự mới tại Đại lý.")

    p_s2_talk = doc.add_paragraph()
    p_s2_talk.add_run("🎙️ Lời thoại người trình bày:\n").bold = True
    r_talk2 = p_s2_talk.add_run(
        '"Thưa Ban Giám khảo, khi một nhân sự mới gia nhập đại lý VinFast — dù là Kế toán, Tư vấn bán hàng hay Kỹ thuật viên — '
        "họ đều phải đối mặt với 3 thách thức lớn:\n"
        "1. Tài liệu nghiệp vụ đồ sộ và phân tán: Hàng trăm tệp DOCX, PDF, bảng giá, slide đào tạo và video kỹ thuật dài hàng giờ.\n"
        "2. Thời gian Onboarding kéo dài: Mất từ 3 đến 4 tuần để nhân viên mới nắm bắt được quy trình hệ thống DMS và chính sách bán hàng.\n"
        "3. Rủi ro sai sót tác nghiệp: Nhầm lẫn trong thủ tục thu hồi pin, xuất hóa đơn hoặc tư vấn sai chính sách ưu đãi làm ảnh hưởng uy tín thương hiệu.\n"
        'Đó chính là động lực để The Sigmoid phát triển giải pháp toàn diện từ hạ tầng dữ liệu Cloud-Native đến AI tác nghiệp thông minh."'
    )
    r_talk2.font.italic = True

    # Slide 3
    h2_3 = doc.add_heading(level=2)
    h2_3.add_run("Slide 3: Kiến trúc Giải pháp Tổng thể (Thời lượng: 1.5 phút)")
    p_s3_act = doc.add_paragraph()
    p_s3_act.add_run("🎬 Thao tác màn hình: ").bold = True
    p_s3_act.add_run("Chiếu Sơ đồ Kiến trúc 2 Track (Data Ingestion & Multi-Agent RAG).")

    p_s3_talk = doc.add_paragraph()
    p_s3_talk.add_run("🎙️ Lời thoại người trình bày:\n").bold = True
    r_talk3 = p_s3_talk.add_run(
        '"Hệ thống của chúng em được kiến trúc theo chuẩn Cloud-Native gồm 2 cấu phần chính:\n\n'
        "• Track 1 — Pipeline Xử lý Dữ liệu Đa phương tiện:\n"
        "  - Tự động thu nạp đa định dạng từ MinIO / AWS S3, chuyển đổi âm thanh video qua Whisper STT, bóc tách bảng biểu và lọc sạch dữ liệu nhạy cảm (PII Masking).\n"
        "  - Lập chỉ mục bằng cơ chế Hybrid Search: kết hợp Dense Vector (ChromaDB) và Sparse Keyword (BM25), phối hợp mô hình Cross-Encoder Reranker để đạt độ chính xác tối đa.\n\n"
        "• Track 2 — Nền tảng Onboarding & Multi-Agent RAG:\n"
        "  - Điều phối luồng bởi LangGraph Agent với bộ phân loại ý định (Intent Routing).\n"
        "  - Áp dụng cơ chế Phân quyền vai trò RBAC nghiêm ngặt: Kế toán, Bán hàng, Kỹ thuật viên chỉ được tiếp cận đúng tài liệu của phòng ban mình, ngăn chặn 100% rò rỉ dữ liệu.\n\n"
        'Ngay sau đây, em xin chuyển sang phần LIVE DEMO thực tế trên hệ thống để Ban Giám khảo cùng trải nghiệm!"'
    )
    r_talk3.font.italic = True

    doc.add_page_break()

    # Section 2: Live Demo End-to-End User Flow
    h1_2 = doc.add_heading(level=1)
    r_h1_2 = h1_2.add_run("II. KỊCH BẢN LIVE DEMO USER FLOW END-TO-END (5 PHÚT)")
    r_h1_2.font.color.rgb = COLOR_PRIMARY

    # Accounts Table
    p_tbl_lbl = doc.add_paragraph()
    p_tbl_lbl.add_run("📋 BẢNG TÀI KHOẢN SỬ DỤNG TRONG PHIÊN DEMO").bold = True
    p_tbl_lbl.runs[0].font.color.rgb = COLOR_PRIMARY

    table_acc = doc.add_table(rows=5, cols=4)
    table_acc.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Vai trò (Role)", "Email Đăng nhập", "Mật khẩu", "Trọng tâm Demo"]
    for i, h in enumerate(headers):
        cell = table_acc.rows[0].cells[i]
        set_cell_background(cell, "0F6B3A")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_acc = [
        (
            "Tư vấn Bán hàng (Sales)",
            "sales@vinfast.vn",
            "123456",
            "Lộ trình Sales, Video S3, Bảng giá xe, Ưu đãi voucher 600k",
        ),
        (
            "Kỹ thuật viên (Technician)",
            "kythuat@vinfast.vn",
            "123456",
            "Test RBAC chặn tài liệu Sales, tra cứu lỗi & bảo dưỡng pin",
        ),
        ("Kế toán viên (Accountant)", "ketoan@vinfast.vn", "123456", "Quy trình DMS, hóa đơn & hợp đồng thuê pin"),
        (
            "Chủ đại lý (Owner)",
            "thehung@vinfast.vn",
            "12345678",
            "Dashboard tiến độ nhân viên toàn đại lý, duyệt Ticket hỗ trợ",
        ),
    ]
    for row_idx, row_data in enumerate(data_acc, start=1):
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = table_acc.rows[row_idx].cells[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p = cell.paragraphs[0]
            p.add_run(text)

    doc.add_paragraph()

    # Step 1
    h3_1 = doc.add_heading(level=2)
    h3_1.add_run("Bước 1: Đăng nhập & Lộ trình Học tập Sales (1 Phút)")
    p_step1 = doc.add_paragraph()
    p_step1.add_run("1. Thao tác: ").bold = True
    p_step1.add_run("Đăng nhập bằng tài khoản ")
    p_step1.add_run("sales@vinfast.vn / 123456").bold = True
    p_step1.add_run(".\n2. ")
    p_step1.add_run("Trải nghiệm Đa phương tiện: ").bold = True
    p_step1.add_run("Bấm vào bài học 'Tổng quan các dòng xe máy điện VinFast'. Bấm nút 'Xem video' / 'Xem tài liệu'.\n")
    p_step1.add_run("3. ")
    p_step1.add_run("Điểm nhấn kỹ thuật: ").bold = True
    p_step1.add_run(
        "Video MP4 (54.4 MB) và tài liệu PDF được stream trực tiếp từ Cloud S3 Backend Proxy mượt mà không cần tải file về máy."
    )

    # Step 2
    h3_2 = doc.add_heading(level=2)
    h3_2.add_run("Bước 2: Trắc nghiệm Tình huống & Tốt nghiệp Onboarding (1 Phút)")
    p_step2 = doc.add_paragraph()
    p_step2.add_run("1. Thao tác: ").bold = True
    p_step2.add_run("Trong bài học, kéo xuống bấm nút xanh ")
    p_step2.add_run("'Trắc nghiệm tình huống'").bold = True
    p_step2.add_run(". Chọn đáp án và bấm nộp bài.\n")
    p_step2.add_run("2. ")
    p_step2.add_run("Tính năng nổi bật: ").bold = True
    p_step2.add_run(
        "Hệ thống đưa ra lời giải thích chi tiết vì sao đúng/sai theo đúng chuẩn quy trình tác nghiệp VinFast. Tiến độ hoàn thành tự động tăng động.\n"
    )
    p_step2.add_run("3. ")
    p_step2.add_run("Chứng nhận: ").bold = True
    p_step2.add_run("Khi đạt 100%, mở thẻ Bài thi Tốt nghiệp để nhận danh hiệu ")
    p_step2.add_run("'🏆 Chứng nhận Tốt nghiệp Onboarding - Đã đạt'").bold = True
    p_step2.add_run(".")

    # Step 3
    h3_3 = doc.add_heading(level=2)
    h3_3.add_run("Bước 3: Trợ lý AI RAG Đa tác vụ & Bảng biểu Tự động (1.5 Phút)")
    p_step3 = doc.add_paragraph()
    p_step3.add_run("1. Mở cửa sổ VF AI Assistant (góc phải màn hình).\n")
    p_step3.add_run("2. Câu hỏi 1 (Bảng giá & Chính sách): ").bold = True
    p_step3.add_run("Gõ: ")
    p_step3.add_run('"chính sách bán hàng và bảng giá xe máy điện"').bold = True
    p_step3.add_run("\n   ➡️ ")
    p_step3.add_run("Kết quả AI: ").bold = True
    p_step3.add_run(
        "AI tự động vẽ bảng giá chi tiết kèm pin/thuê pin của toàn bộ xe (Feliz II: 28.7M/23.0M, Evo, Viper, Vento S...), voucher chăm sóc xe 600k và ưu đãi 2% trước bạ kèm trích dẫn nguồn '260801 Chính sách bán hàng XMĐ'.\n"
    )
    p_step3.add_run("3. Câu hỏi 2 (Kỹ năng tác nghiệp): ").bold = True
    p_step3.add_run("Gõ: ")
    p_step3.add_run('"kỹ thuật bán hàng"').bold = True
    p_step3.add_run("\n   ➡️ ")
    p_step3.add_run("Kết quả AI: ").bold = True
    p_step3.add_run(
        "AI trích xuất quy trình tư vấn 7 bước, phương pháp chốt hợp đồng 1-2-3 và kỹ thuật xử lý từ chối FABE & LACE."
    )

    # Step 4
    h3_4 = doc.add_heading(level=2)
    h3_4.add_run("Bước 4: Trình diễn Phân quyền Bảo mật RBAC & Escalation (1.5 Phút)")
    p_step4 = doc.add_paragraph()
    p_step4.add_run("1. Thao tác: ").bold = True
    p_step4.add_run("Đăng xuất ➡️ Đăng nhập tài khoản Kỹ thuật viên ")
    p_step4.add_run("kythuat@vinfast.vn / 123456").bold = True
    p_step4.add_run(".\n2. ")
    p_step4.add_run("Test chặn quyền: ").bold = True
    p_step4.add_run("Hỏi câu hỏi Sales: ")
    p_step4.add_run('"Quy trình 7 bước tư vấn bán hàng"').bold = True
    p_step4.add_run("\n   ➡️ ")
    p_step4.add_run("Kết quả AI: ").bold = True
    p_step4.add_run(
        "AI ngay lập tức từ chối cung cấp: 'Không tìm thấy thông tin phù hợp trong tài liệu được cấp quyền' (Do tài liệu thuộc role: sales).\n"
    )
    p_step4.add_run("3. ")
    p_step4.add_run("Test Escalation Ticket: ").bold = True
    p_step4.add_run("Gõ sự cố: ")
    p_step4.add_run('"Khách hàng yêu cầu đổi xe mới ngoài chính sách bảo hành"').bold = True
    p_step4.add_run("\n   ➡️ ")
    p_step4.add_run("Kết quả AI: ").bold = True
    p_step4.add_run(
        "AI nhận diện tình huống cần thẩm quyền giải quyết, tự động gắn nhãn [Chuyển tiếp IT/Manager] và hiện nút "
    )
    p_step4.add_run("'Tạo phiếu hỗ trợ (Support Ticket)'").bold = True
    p_step4.add_run(" gửi thẳng lên Chủ đại lý.\n")
    p_step4.add_run("4. ")
    p_step4.add_run("Đăng nhập Chủ đại lý (thehung@vinfast.vn): ").bold = True
    p_step4.add_run("Xem Dashboard theo dõi tiến độ toàn đại lý và phê duyệt Ticket hỗ trợ.")

    doc.add_page_break()

    # Section 3: Q&A Cheatsheet
    h1_3 = doc.add_heading(level=1)
    r_h1_3 = h1_3.add_run("III. BỘ CÂU HỎI PHẢN BIỆN THƯỜNG GẶP (Q&A CHEATSHEET)")
    r_h1_3.font.color.rgb = COLOR_PRIMARY

    qa_list = [
        (
            "Câu 1: Làm sao hệ thống đảm bảo AI không bị 'ảo giác' (Hallucination) khi trả lời chính sách giá?",
            "Trả lời: Hệ thống áp dụng Strict RAG Context Grounding. System Prompt bắt buộc AI chỉ được trả lời dựa trên các đoạn văn bản trích xuất từ ChromaDB & BM25 và phải trích dẫn rõ [TÀI LIỆU i]. Nếu ngữ cảnh không có thông tin, AI buộc phải trả về thông báo từ chối thay vì tự suy đoán. Đồng thời mô hình Cross-Encoder Reranker đảm bảo chọn đúng top chunk có liên quan cao nhất.",
        ),
        (
            "Câu 2: Cơ chế Phân quyền RBAC hoạt động ở tầng nào? Có an toàn không?",
            "Trả lời: RBAC được thực thi ở 2 tầng độc lập: (1) Tầng API: FastAPI Dependencies kiểm tra JWT role trước khi phục vụ file/video từ S3. (2) Tầng Vector Store: ChromaDB lọc metadata `role: {'$in': access_scope}` trước khi nạp vào LLM. Kể cả người dùng có cố tình prompt injection thì LLM cũng không có dữ liệu trong ngữ cảnh để đọc.",
        ),
        (
            "Câu 3: Nếu dữ liệu có thêm tài liệu mới tải lên thì hệ thống có phải đào tạo lại model không?",
            "Trả lời: Hoàn toàn không! Đây là ưu thế tuyệt đối của kiến trúc RAG. Khi Admin tải file mới lên AWS S3 / MinIO, Pipeline Cloud Ingestion tự động bóc tách, embed và nạp vào ChromaDB + BM25 trong vài giây. Chatbot có thể tra cứu ngay kiến thức mới mà không cần re-train hoặc fine-tune LLM.",
        ),
        (
            "Câu 4: Hệ thống xử lý các file đa phương tiện (Video/Âm thanh) như thế nào?",
            "Trả lời: Trong Track 1, chúng em sử dụng mô hình OpenAI Whisper STT để nhận dạng tiếng Việt có gắn mốc thời gian (timestamps). Văn bản sau đó được lọc PII, ghép đoạn ngữ nghĩa (Semantic Chunking) và đánh chỉ mục song song cùng các tài liệu PDF/DOCX.",
        ),
    ]

    for q, a in qa_list:
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(f"❓ {q}\n")
        r_q.bold = True
        r_q.font.color.rgb = COLOR_PRIMARY
        r_q.font.size = Pt(11)

        p_a = doc.add_paragraph()
        r_a = p_a.add_run(f"💡 {a}\n")
        r_a.font.size = Pt(10.5)

    # Save docx
    output_path = (
        "/Users/sethehung/Documents/aithucchien_vinuni/Project/team-The_sigmoid/VF_AI_Presentation_and_Demo_Script.docx"
    )
    doc.save(output_path)
    print(f"Successfully generated DOCX file at: {output_path}")


if __name__ == "__main__":
    create_document()
